from __future__ import annotations

import math
import shutil
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from fastapi import UploadFile
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.document import Document, DocumentChunk
from app.models.settings import UserSetting
from app.providers.registry import DefaultProviderRegistry


class DocumentService:
    def __init__(self) -> None:
        self.settings = get_settings()

    def list_documents(self, db: Session, user_id: str) -> list[Document]:
        return list(db.scalars(select(Document).where(Document.user_id == user_id).order_by(Document.created_at.desc())))

    def _build_storage_path(self, original_name: str) -> Path:
        suffix = Path(original_name).suffix
        return self.settings.storage_dir / "documents" / f"{uuid4()}{suffix}"

    def _read_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return file_path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".pdf":
            reader = PdfReader(str(file_path))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        if suffix == ".docx":
            document = DocxDocument(str(file_path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        return file_path.read_text(encoding="utf-8", errors="ignore")

    def _chunk_text(self, text: str, chunk_size: int = 850, overlap: int = 120) -> list[str]:
        normalized = " ".join(text.split())
        if not normalized:
            return []
        chunks: list[str] = []
        start = 0
        while start < len(normalized):
            end = min(len(normalized), start + chunk_size)
            chunks.append(normalized[start:end])
            if end == len(normalized):
                break
            start = max(end - overlap, start + 1)
        return chunks

    def _cosine_similarity(self, left: list[float], right: list[float]) -> float:
        numerator = sum(a * b for a, b in zip(left, right, strict=False))
        left_norm = math.sqrt(sum(a * a for a in left))
        right_norm = math.sqrt(sum(b * b for b in right))
        if left_norm == 0 or right_norm == 0:
            return 0.0
        return numerator / (left_norm * right_norm)

    async def ingest_upload(
        self,
        db: Session,
        *,
        user_id: str,
        upload_file: UploadFile,
        providers: DefaultProviderRegistry,
        provider_settings: UserSetting,
    ) -> Document:
        storage_path = self._build_storage_path(upload_file.filename or "document.txt")
        document = Document(
            user_id=user_id,
            filename=storage_path.name,
            original_name=upload_file.filename or storage_path.name,
            file_path=str(storage_path),
            file_size_bytes=None,
            mime_type=upload_file.content_type,
            status="processing",
        )
        db.add(document)
        db.commit()
        db.refresh(document)

        try:
            with storage_path.open("wb") as target:
                shutil.copyfileobj(upload_file.file, target)

            document.file_size_bytes = storage_path.stat().st_size
            text = self._read_text(storage_path)
            chunks = self._chunk_text(text)
            embeddings: list[list[float] | None]
            if chunks:
                try:
                    provider = providers.get_embedding_provider(provider_settings.embedding_provider, provider_settings)
                    embeddings = await provider.embed(
                        texts=chunks,
                        model=provider_settings.embedding_model,
                        dimensions=1536,
                    )
                except Exception:
                    embeddings = [None for _ in chunks]
            else:
                embeddings = []

            for existing_chunk in list(document.chunks):
                db.delete(existing_chunk)

            for index, chunk in enumerate(chunks):
                db.add(
                    DocumentChunk(
                        document_id=document.id,
                        chunk_index=index,
                        content=chunk,
                        embedding=embeddings[index] if index < len(embeddings) else None,
                    )
                )

            document.status = "ready"
            document.error_message = None
            db.add(document)
            db.commit()
            db.refresh(document)
            return document
        except Exception as exc:
            document.status = "error"
            document.error_message = str(exc)
            db.add(document)
            db.commit()
            db.refresh(document)
            return document

    def delete_document(self, db: Session, document_id: str) -> None:
        document = db.scalar(select(Document).where(Document.id == document_id))
        if document is None:
            raise ValueError("Document not found")
        if document.file_path:
            path = Path(document.file_path)
            if path.exists():
                path.unlink()
        db.delete(document)
        db.commit()

    async def similarity_context(
        self,
        db: Session,
        *,
        transcript_text: str,
        providers: DefaultProviderRegistry | None = None,
        provider_settings: UserSetting | None = None,
        limit: int = 3,
    ) -> str:
        chunks = list(db.scalars(select(DocumentChunk).order_by(DocumentChunk.created_at.desc())))
        if providers and provider_settings and transcript_text.strip():
            try:
                provider = providers.get_embedding_provider(provider_settings.embedding_provider, provider_settings)
                query_embedding = (await provider.embed(
                    texts=[transcript_text],
                    model=provider_settings.embedding_model,
                    dimensions=1536,
                ))[0]
                ranked = [
                    chunk
                    for chunk in sorted(
                        [chunk for chunk in chunks if chunk.embedding],
                        key=lambda item: self._cosine_similarity(item.embedding, query_embedding),
                        reverse=True,
                    )[:limit]
                ]
                if ranked:
                    return "\n\n".join(chunk.content for chunk in ranked)
            except Exception:
                pass

        return "\n\n".join(chunk.content for chunk in chunks[:limit])
